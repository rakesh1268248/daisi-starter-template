import PyPDF2 as pdf
import docx
from pptx import Presentation
import re
import string
import nltk
import streamlit as st
nltk.download('all')
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
stop_words=set(nltk.corpus.stopwords.words('english'))
lemmatizer = WordNetLemmatizer()
from rank_bm25 import *
st.text('debug point1')

def preprocessing(documents):
  documents_clean = []
  for d in documents:
    # Remove Unicode
    document_test = re.sub('[^a-zA-Z0-9]', ' ', str(d))
    # Remove Mentions
    document_test = re.sub(r'@\w+', '', document_test)
    # Lowercase the document
    document_test = document_test.lower()
    # Remove punctuations
    document_test = re.sub(r'[%s]' % re.escape(string.punctuation), ' ', document_test)
    # Lowercase the numbers
    document_test = re.sub(r'[0-9]', '', document_test)
    # Remove the doubled space
    document_test = re.sub(r'\s{2,}', ' ', document_test)
    #tokenization
    document_test = document_test.split()
    #lemmmitization
    document_test = [lemmatizer.lemmatize(word) for word in document_test if not word in set(stopwords.words('english'))]
    document_test = ' '.join(document_test)
    documents_clean.append(document_test)
  return documents_clean

def search_report(documents_clean,Enter_text):
  tokenized_corpus = [doc.split(" ") for doc in documents_clean]
  bm25 = BM25Okapi(tokenized_corpus)
  tokenized_query = Enter_text.split()
  # doc_scores = bm25.get_scores(tokenized_query)
  result=bm25.get_top_n(tokenized_query,documents_clean , n=5)
  st.text(result)
  
def st_ui():
  st.set_page_config(layout = "wide")
  st.title("Auto Review Legal contracts - DocumentAI")
    
  fileupload = st.sidebar.file_uploader("Upload a Contract here")
  select_category = st.sidebar.selectbox("Endpoint selection", ["category", "PDF", 'Word Document','PPT'])
  Enter_text = st.sidebar.text_input("Text to search", "please enter the text")
  st.text(fileupload)
  Button=st.sidebar.button('Analyze_contract')
  st.text('debug point 3')
   
  if fileupload is not None:
    text=[]
    st.text('debug point 4')
    if select_category == "PDF":
      f=open(fileupload,'rb')
      reader = pdf.PdfFileReader(f)
      n_pages = len(reader.pages)
      text.append(pageObj.extractText().lower().split('\n'))
    elif select_category =="Word Document":
      doc = docx.Document(fileupload)
      for i in range(len(doc.paragraphs)):
        text.append(doc.paragraphs[i].text)
  cleaned_document=preprocessing(text)
  search_report(cleaned_document)
    

if __name__ == "__main__":
    st_ui()
